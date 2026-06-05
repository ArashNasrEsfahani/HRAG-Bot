"""Document loaders: read files from disk into Document dataclasses.

Dispatch by file extension:
  .pdf      -> pymupdf (fitz) by default; or Nougat-OCR when
               ``cfg.ingest.use_nougat=True`` and nougat-ocr is installed.
  .docx     -> python-docx
  .md/.markdown/.txt -> raw text read

Math note: PyMuPDF cannot reconstruct LaTeX source from PDF math glyphs;
equations come back as unicode (e.g. "α²+β²=γ²") or partial text. The
equation-aware chunker (``ingest.math_detect`` + ``ingest.chunker``) only
detects ``$$...$$`` and ``\\begin{X}...\\end{X}`` blocks, so its boundary
protection mainly helps Markdown, .tex, and HTML-scraped sources. PDFs
remain best-effort: whatever PyMuPDF extracted is what the chunker sees.
Enable ``ingest.use_nougat: true`` (requires ``pip install nougat-ocr``) to
get LaTeX-reconstructed math from academic PDFs.
"""

from __future__ import annotations

import hashlib
import logging
import re
from pathlib import Path
from typing import TYPE_CHECKING, Any

from hrag.types import Document

if TYPE_CHECKING:  # pragma: no cover
    from hrag.config import Config

logger = logging.getLogger(__name__)


# Page-number-only line: matches the ENTIRE line (e.g. "12", "Page 3 of 50").
# Mirrors quality._RE_PAGE_ARTIFACT — kept in sync; intentionally duplicated
# rather than imported to avoid the loader depending on quality.py.
_RE_PAGE_NUMBER_LINE = re.compile(
    r"^(?:Page\s*)?\d+(?:\s*of\s*\d+)?\s*$",
    re.IGNORECASE,
)


def _strip_page_artifacts(page_text: str) -> str:
    """Remove page-number-only lines from a single PDF page's extracted text.

    PyMuPDF returns page text with page numbers/footers as bare lines that
    later get glued onto adjacent content if pages are joined with a single
    newline. Drop those lines here; downstream paragraph splitting then sees
    a clean run of body text.
    """
    kept = [ln for ln in page_text.splitlines() if not _RE_PAGE_NUMBER_LINE.match(ln.strip())]
    return "\n".join(kept)


# ---------------------------------------------------------------------------
# Stable doc_id: SHA-1 of (user_id + absolute path), take first 16 hex chars
# ---------------------------------------------------------------------------

def _make_doc_id(user_id: str, abs_path: str) -> str:
    raw = (user_id + abs_path).encode("utf-8")
    return hashlib.sha1(raw).hexdigest()[:16]


# ---------------------------------------------------------------------------
# Format-specific loaders
# ---------------------------------------------------------------------------

def _load_pdf_pymupdf(path: Path) -> tuple[str, dict[str, Any]]:
    """Load a PDF with PyMuPDF (fitz); return (text, metadata)."""
    try:
        import fitz  # pymupdf
    except ImportError as exc:
        raise ImportError("pymupdf is required for PDF loading: pip install pymupdf") from exc

    doc = fitz.open(str(path))
    pages: list[str] = []
    page_numbers: list[int] = []
    for page in doc:
        text = _strip_page_artifacts(page.get_text("text"))
        if text.strip():
            pages.append(text)
            page_numbers.append(page.number + 1)  # 1-based
    doc.close()

    # Join pages with a blank line so the paragraph splitter (which splits on
    # \n{2,}) treats page boundaries as paragraph boundaries. Without this,
    # the last paragraph of page N glues onto the first paragraph of page N+1.
    # Compute page_spans (half-open char ranges into full_text) at join time —
    # offsets are free here and required by the page-metadata ingest path.
    SEP = "\n\n"
    page_spans: list[tuple[int, int, int]] = []
    cursor = 0
    for text, pno in zip(pages, page_numbers):
        start = cursor
        cursor += len(text)
        page_spans.append((start, cursor, pno))
        cursor += len(SEP)
    full_text = SEP.join(pages)

    metadata: dict[str, Any] = {
        "format": "pdf",
        "page_count": len(page_numbers),
        "pages": page_numbers,
        "page_spans": page_spans,
    }
    return full_text, metadata


def _load_pdf(path: Path, *, cfg: Config | None = None) -> tuple[str, dict[str, Any]]:
    """Load a PDF; dispatch to Nougat-OCR when configured, otherwise PyMuPDF.

    Parameters
    ----------
    path:
        Path to the PDF file.
    cfg:
        Optional top-level :class:`~hrag.config.Config` instance.  When
        ``cfg.ingest.use_nougat`` is ``True`` *and* ``nougat-ocr`` is
        installed, Nougat-OCR is used to extract LaTeX-rich text.  Any
        Nougat failure falls back to PyMuPDF automatically.  When ``cfg`` is
        ``None``, PyMuPDF is used unconditionally.
    """
    use_nougat = bool(cfg and cfg.ingest.use_nougat)
    if use_nougat:
        try:
            from hrag.ingest.nougat_loader import is_nougat_available, load_pdf_nougat

            if is_nougat_available():
                try:
                    text = load_pdf_nougat(
                        path,
                        model=cfg.ingest.nougat_model,  # type: ignore[union-attr]
                    )
                    metadata: dict[str, Any] = {
                        "format": "pdf",
                        "loader": "nougat",
                    }
                    return text, metadata
                except Exception as exc:
                    logger.warning(
                        "Nougat OCR failed for %s; falling back to PyMuPDF: %s",
                        path,
                        exc,
                    )
            else:
                logger.info("Nougat not installed; using PyMuPDF.")
        except ImportError:
            logger.info("Nougat not installed; using PyMuPDF.")

    return _load_pdf_pymupdf(path)


def _load_docx(path: Path) -> tuple[str, dict[str, Any]]:
    """Load a .docx file with python-docx; return (text, metadata)."""
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for .docx loading: pip install python-docx"
        ) from exc

    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs]
    full_text = "\n".join(paragraphs)
    metadata: dict[str, Any] = {
        "format": "docx",
        "paragraph_count": len(paragraphs),
    }
    return full_text, metadata


def _load_text(path: Path, fmt: str) -> tuple[str, dict[str, Any]]:
    """Load a plain-text or markdown file."""
    text = path.read_text(encoding="utf-8", errors="replace")
    metadata: dict[str, Any] = {"format": fmt}
    return text, metadata


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def load_document(
    path: str | Path,
    user_id: str,
    *,
    cfg: Config | None = None,
) -> Document:
    """Load a file from disk and return a Document.

    Supported extensions: .pdf, .docx, .md, .markdown, .txt
    Raises ValueError for unsupported extensions.

    Parameters
    ----------
    path:
        Path to the file on disk.
    user_id:
        ID of the owning user; used to derive a stable ``doc_id``.
    cfg:
        Optional top-level config.  When supplied and
        ``cfg.ingest.use_nougat`` is ``True``, PDF files are loaded with
        Nougat-OCR instead of PyMuPDF (falls back automatically when the
        dep is absent or Nougat raises).
    """
    path = Path(path).resolve()
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    ext = path.suffix.lower()
    if ext == ".pdf":
        text, metadata = _load_pdf(path, cfg=cfg)
    elif ext == ".docx":
        text, metadata = _load_docx(path)
    elif ext in (".md", ".markdown"):
        text, metadata = _load_text(path, "md")
    elif ext == ".txt":
        text, metadata = _load_text(path, "txt")
    else:
        raise ValueError(
            f"Unsupported file extension {ext!r}. "
            "Supported: .pdf, .docx, .md, .markdown, .txt"
        )

    doc_id = _make_doc_id(user_id, str(path))

    return Document(
        doc_id=doc_id,
        user_id=user_id,
        source_path=str(path),
        title=path.stem,
        text=text,
        source_type="document",
        metadata=metadata,
    )
